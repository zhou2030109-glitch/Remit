"""Paper post-processing plus reproducible LaTeX/PDF delivery."""

from __future__ import annotations

import hashlib
import json
import math
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps

from app.config.setting import settings
from app.schemas.enums import CompTemplate
from app.utils.log_util import logger

MAX_ABSTRACT_CHARS = 700
IMAGE_BLOCK_RE = re.compile(r"^\s*!\[(?P<alt>.*?)\]\((?P<src>.*?)\)\s*$")
HEADING_RE = re.compile(r"^\s*#{1,6}\s+")
ABSTRACT_RE = re.compile(r"^\s*#{1,6}\s*摘要\s*$")
KEYWORD_RE = re.compile(r"^\s*\*{0,2}\s*关(?:键)?词[:：]?\s*\*{0,2}")
QUESTION_LEAD_RE = re.compile(
    r"^(?P<prefix>针对)?(?P<lead>问题[一二三四五六七八九十])[,，：: ]*(?P<body>.*)$",
)

_LEGACY_PAPER_OUTPUTS = (
    "res.md",
    "res_polished.md",
    "res.docx",
    "res_polished.docx",
    "res_polished.pdf",
    "paper_render.html",
    "paper_pdf_header.tex",
    "paper_reference.docx",
)


class PaperRenderError(RuntimeError):
    """LaTeX 生成、编译或 PDF 复核失败。"""


@dataclass(frozen=True)
class PaperDeliverables:
    """由同一份可编译 LaTeX 源码生成的最终交付物。"""

    tex_path: Path
    pdf_path: Path
    report_path: Path
    page_count: int


def render_paper_deliverables(
    markdown: str,
    work_dir: str | Path,
    comp_template: CompTemplate,
) -> PaperDeliverables:
    """将已校验终稿转换为 LaTeX，并由该源码编译、复核 PDF。"""
    root = Path(work_dir).resolve()
    root.mkdir(parents=True, exist_ok=True)
    build_dir = root / ".remit" / "latex-build"
    build_dir.mkdir(parents=True, exist_ok=True)
    tex_path = root / "res.tex"
    pdf_path = root / "res.pdf"
    report_path = root / "paper_delivery_report.json"
    pdf_path.unlink(missing_ok=True)
    report_path.unlink(missing_ok=True)

    _convert_markdown_to_latex(markdown, tex_path, root, build_dir, comp_template)
    _validate_latex_source(tex_path)
    engine = _compile_latex(tex_path, build_dir)
    built_pdf = build_dir / "res.pdf"
    if not built_pdf.is_file():
        raise PaperRenderError("LaTeX 编译命令成功退出，但没有生成 res.pdf")
    shutil.copy2(built_pdf, pdf_path)
    pdf_metrics = inspect_pdf_artifact(
        pdf_path,
        comp_template=comp_template,
        minimum_pages=settings.PAPER_MIN_PDF_PAGES,
    )
    report = {
        "status": "pass",
        "source": tex_path.name,
        "pdf": pdf_path.name,
        "compiler": engine,
        "compile_passes": 2,
        "tex_sha256": _sha256(tex_path),
        "pdf_sha256": _sha256(pdf_path),
        **pdf_metrics,
    }
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    _remove_legacy_paper_outputs(root)
    return PaperDeliverables(
        tex_path=tex_path,
        pdf_path=pdf_path,
        report_path=report_path,
        page_count=int(pdf_metrics["page_count"]),
    )


def _convert_markdown_to_latex(
    markdown: str,
    output_path: Path,
    resource_path: Path,
    build_dir: Path,
    comp_template: CompTemplate,
) -> None:
    import pypandoc  # type: ignore[import-unresolved]

    header_path = build_dir / "paper_header.tex"
    header_path.write_text(
        build_pdf_header(resource_path, comp_template), encoding="utf-8"
    )
    try:
        pypandoc.convert_text(
            markdown,
            to="latex",
            format="markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables+raw_html",
            outputfile=str(output_path),
            extra_args=[
                f"--resource-path={resource_path}",
                f"--include-in-header={header_path}",
                "--standalone",
                "--wrap=none",
            ],
        )
    except Exception as exc:
        raise PaperRenderError(f"Pandoc 生成 LaTeX 失败: {exc}") from exc


def _validate_latex_source(tex_path: Path) -> None:
    if not tex_path.is_file() or tex_path.stat().st_size < 500:
        raise PaperRenderError("生成的 res.tex 为空或内容异常短")
    source = tex_path.read_text(encoding="utf-8")
    required = ("\\documentclass", "\\begin{document}", "\\end{document}")
    missing = [token for token in required if token not in source]
    if missing:
        raise PaperRenderError("res.tex 缺少完整文档结构: " + ", ".join(missing))
    markdown_leaks = re.findall(r"(?m)^\s*(?:#{1,6}\s+|!\[[^\]]*\]\([^)]+\))", source)
    if markdown_leaks:
        raise PaperRenderError("res.tex 仍包含未转换的 Markdown 块标记")


def _compile_latex(tex_path: Path, build_dir: Path) -> str:
    configured = settings.LATEX_ENGINE.strip() or "xelatex"
    engine = shutil.which(configured)
    if not engine:
        raise PaperRenderError(
            f"找不到 LaTeX 编译器 {configured}；请安装 MiKTeX 或 TeX Live 并加入 PATH"
        )

    command = [
        engine,
        "-no-shell-escape",
        "-interaction=nonstopmode",
        "-halt-on-error",
        "-file-line-error",
        f"-output-directory={build_dir}",
        tex_path.name,
    ]
    logs: list[str] = []
    for compile_pass in range(1, 3):
        try:
            completed = subprocess.run(
                command,
                cwd=tex_path.parent,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=settings.LATEX_COMPILE_TIMEOUT_SECONDS,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise PaperRenderError(
                f"LaTeX 第 {compile_pass} 次编译超过 "
                f"{settings.LATEX_COMPILE_TIMEOUT_SECONDS:g} 秒"
            ) from exc
        output = (completed.stdout or "") + "\n" + (completed.stderr or "")
        logs.append(f"===== pass {compile_pass} =====\n{output}")
        if completed.returncode != 0:
            (build_dir / "compile-output.txt").write_text(
                "\n".join(logs), encoding="utf-8"
            )
            tail = output[-3000:].strip()
            raise PaperRenderError(
                f"res.tex 第 {compile_pass} 次编译失败（退出码 "
                f"{completed.returncode}）：{tail}"
            )
    (build_dir / "compile-output.txt").write_text(
        "\n".join(logs), encoding="utf-8"
    )
    return Path(engine).name


def inspect_pdf_artifact(
    pdf_path: str | Path,
    *,
    comp_template: CompTemplate,
    minimum_pages: int,
) -> dict[str, object]:
    """重开并渲染抽样页面，拒绝空白、损坏或纸型错误的 PDF。"""
    import pymupdf

    path = Path(pdf_path)
    if not path.is_file() or path.stat().st_size < 1_000:
        raise PaperRenderError("res.pdf 缺失或内容异常短")
    try:
        document = pymupdf.open(path)
    except Exception as exc:
        raise PaperRenderError(f"res.pdf 无法重新打开: {exc}") from exc

    try:
        if document.needs_pass:
            raise PaperRenderError("res.pdf 被意外加密，无法验收")
        page_count = document.page_count
        if page_count < minimum_pages:
            raise PaperRenderError(
                f"res.pdf 仅 {page_count} 页，少于终稿下限 {minimum_pages} 页"
            )
        first_rect = document[0].rect
        expected = (
            (595.3, 841.9)
            if comp_template == CompTemplate.CHINA
            else (612.0, 792.0)
        )
        if abs(first_rect.width - expected[0]) > 8 or abs(first_rect.height - expected[1]) > 8:
            label = "A4" if comp_template == CompTemplate.CHINA else "US Letter"
            raise PaperRenderError(f"res.pdf 首页纸型不是要求的 {label}")

        blank_pages: list[int] = []
        total_text_chars = 0
        for index, page in enumerate(document):
            page_text = page.get_text().strip()
            total_text_chars += len(page_text)
            if not page_text and not page.get_images() and not page.get_drawings():
                blank_pages.append(index + 1)
            for block in page.get_text("blocks"):
                x0, y0, x1, y1 = block[:4]
                if x0 < -2 or y0 < -2 or x1 > page.rect.width + 2 or y1 > page.rect.height + 2:
                    raise PaperRenderError(f"res.pdf 第 {index + 1} 页存在越界文本")
        if blank_pages:
            raise PaperRenderError(
                "res.pdf 含完全空白页: " + ", ".join(map(str, blank_pages))
            )
        if total_text_chars < 1_000:
            raise PaperRenderError("res.pdf 可提取正文不足 1000 字符，疑似字体或渲染损坏")

        sampled_pages = sorted({0, page_count // 2, page_count - 1})
        for index in sampled_pages:
            pixmap = document[index].get_pixmap(matrix=pymupdf.Matrix(1, 1), alpha=False)
            sample = pixmap.samples[:: max(1, pixmap.n * 20)]
            if not sample or all(value > 248 for value in sample):
                raise PaperRenderError(f"res.pdf 第 {index + 1} 页渲染结果近似全白")
        return {
            "page_count": page_count,
            "paper_size": "A4" if comp_template == CompTemplate.CHINA else "Letter",
            "blank_pages": blank_pages,
            "extracted_text_chars": total_text_chars,
            "rendered_pages_checked": [index + 1 for index in sampled_pages],
        }
    finally:
        document.close()


def _remove_legacy_paper_outputs(root: Path) -> None:
    for filename in _LEGACY_PAPER_OUTPUTS:
        (root / filename).unlink(missing_ok=True)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def render_paper_docx(task_id: str) -> Path:
    """Render polished DOCX for a completed task."""
    work_dir = Path("project") / "work_dir" / task_id
    work_dir.mkdir(parents=True, exist_ok=True)

    source_md = work_dir / "res.md"
    if not source_md.exists():
        raise FileNotFoundError(f"missing paper markdown: {source_md}")

    polished_md = work_dir / "res_polished.md"
    raw_markdown = source_md.read_text(encoding="utf-8")
    polished_markdown = polish_markdown(raw_markdown, work_dir)
    polished_md.write_text(polished_markdown, encoding="utf-8")

    reference_docx = build_reference_docx(work_dir)
    docx_path = work_dir / "res.docx"
    polished_docx_path = work_dir / "res_polished.docx"

    convert_markdown_to_docx(
        markdown=polished_markdown,
        output_path=docx_path,
        reference_docx=reference_docx,
        resource_path=work_dir,
    )
    shutil.copy2(docx_path, polished_docx_path)
    try:
        convert_markdown_to_pdf(
            markdown=polished_markdown,
            output_path=work_dir / "res_polished.pdf",
            resource_path=work_dir,
        )
    except Exception as exc:  # pragma: no cover - best effort export
        logger.warning("pdf export skipped: %s", exc)
    logger.info("paper docx generated: %s", docx_path)
    logger.info("paper polished docx copied: %s", polished_docx_path)
    return docx_path


def polish_markdown(markdown: str, work_dir: Path) -> str:
    """Apply paper-level Markdown cleanup."""
    markdown = markdown.replace("\r\n", "\n")
    markdown = normalize_common_math(markdown)
    markdown = compact_abstract(markdown)
    markdown = merge_image_blocks(markdown, work_dir)
    markdown = append_source_code_appendix(markdown, work_dir)
    markdown = ensure_blank_lines_around_headings(markdown)
    return markdown.strip() + "\n"


def normalize_common_math(markdown: str) -> str:
    """Normalize a few math-like tokens so pandoc renders them more reliably."""
    replacements = {
        r"(?<![\\$])\bR\^2\b(?![\\$])": r"$R^2$",
        r"(?<![\\$])\bR\^3\b(?![\\$])": r"$R^3$",
        r"(?<![\\$])\bOOF\s+R\^2\b(?![\\$])": r"OOF $R^2$",
    }
    lines = markdown.splitlines()
    out: list[str] = []
    in_display_math = False

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if stripped == "$$":
            out.append(line)
            in_display_math = not in_display_math
            idx += 1
            continue
        if in_display_math:
            if stripped == "R^2" and idx + 1 < len(lines):
                next_stripped = lines[idx + 1].strip()
                if next_stripped.startswith("="):
                    out.append("R^2 " + next_stripped)
                    idx += 2
                    continue
            out.append(line)
            idx += 1
            continue
        for pattern, repl in replacements.items():
            line = re.sub(pattern, repl, line)
        out.append(line)
        idx += 1

    return "\n".join(out)


def compact_abstract(markdown: str) -> str:
    """Shorten the abstract and emphasize lead phrases."""
    lines = markdown.splitlines()
    start = None
    start_level = 0

    for idx, line in enumerate(lines):
        if ABSTRACT_RE.match(line):
            start = idx
            start_level = len(line) - len(line.lstrip("#"))
            break
    if start is None:
        return markdown

    end = len(lines)
    for idx in range(start + 1, len(lines)):
        line = lines[idx]
        if HEADING_RE.match(line):
            level = len(line) - len(line.lstrip("#"))
            if level <= start_level:
                end = idx
                break

    prefix = lines[: start + 1]
    body = lines[start + 1 : end]
    suffix = lines[end:]

    paragraphs = split_paragraphs(body)
    abstract_blocks: list[str] = []
    keyword_lines: list[str] = []

    for paragraph in paragraphs:
        if KEYWORD_RE.match(paragraph):
            keyword_lines.append("**关键词：** " + KEYWORD_RE.sub("", paragraph).strip())
            continue
        abstract_blocks.append(compact_abstract_paragraph(paragraph))

    if len("\n\n".join(abstract_blocks)) > MAX_ABSTRACT_CHARS:
        abstract_blocks = [tighten_paragraph(p) for p in abstract_blocks]
    if len("\n\n".join(abstract_blocks)) > MAX_ABSTRACT_CHARS:
        abstract_blocks = [compress_abstract_paragraph(p) for p in abstract_blocks]
    if len("\n\n".join(abstract_blocks)) > MAX_ABSTRACT_CHARS:
        abstract_blocks = [limit_abstract_paragraph(p, 140 if idx == 0 else 110) for idx, p in enumerate(abstract_blocks)]

    rebuilt: list[str] = [*prefix, ""]
    for block in abstract_blocks:
        if block.strip():
            rebuilt.extend(block.splitlines())
            rebuilt.append("")
    for line in keyword_lines:
        rebuilt.append(line)
    if keyword_lines:
        rebuilt.append("")
    rebuilt.extend(suffix)
    return "\n".join(rebuilt)


def compact_abstract_paragraph(paragraph: str) -> str:
    """Compress a single abstract paragraph."""
    text = paragraph.strip()
    if not text:
        return text

    text = text.replace("**", "")
    match = QUESTION_LEAD_RE.match(text)
    if match:
        lead = (match.group("prefix") or "") + match.group("lead")
        rest = match.group("body").strip(" ,，：:")
        sentences = split_sentences(rest)
        result = next((s for s in sentences if sentence_has_signal(s)), "")
        if not result and sentences:
            result = sentences[0]
        result = squeeze_clause(result)
        if result:
            return f"**{lead}：** **{result}**"
        return f"**{lead}：**"

    sentences = split_sentences(text)
    if len(sentences) <= 2:
        return text

    first = sentences[0]
    second = next((s for s in sentences[1:] if sentence_has_signal(s)), "")
    if not second and len(sentences) > 1:
        second = sentences[1]
    if second:
        return first + second
    return first


def compress_abstract_paragraph(paragraph: str) -> str:
    """Aggressively compress an abstract paragraph when it is still too long."""
    text = paragraph.strip().replace("**", "")
    match = QUESTION_LEAD_RE.match(text)
    if match:
        lead = (match.group("prefix") or "") + match.group("lead")
        rest = match.group("body").strip(" ,，：:")
        sentences = split_sentences(rest)
        result = next((s for s in sentences if sentence_has_signal(s)), "")
        if not result and sentences:
            result = sentences[0]
        result = squeeze_clause(result)
        return f"**{lead}：** {result}" if result else f"**{lead}：**"

    sentences = split_sentences(text)
    if not sentences:
        return text
    signal = next((s for s in sentences if sentence_has_signal(s)), sentences[0])
    return squeeze_clause(signal)


def limit_abstract_paragraph(paragraph: str, max_chars: int) -> str:
    """Final fallback to keep the abstract within one page."""
    text = paragraph.strip()
    if len(text) <= max_chars:
        return text
    if text.startswith("**") and "** " in text:
        lead, body = text.split("** ", 1)
        body = body[: max_chars - len(lead) - 5].rstrip("，,；;。 ") + "…"
        return f"{lead}** {body}"
    return text[: max_chars - 1].rstrip("，,；;。 ") + "…"


def squeeze_clause(text: str) -> str:
    """Trim a sentence to its most informative clause."""
    text = text.strip()
    if not text:
        return text
    for delimiter in ("；", ";"):
        if delimiter in text:
            parts = [part.strip() for part in text.split(delimiter) if part.strip()]
            if parts:
                return parts[0]
    return text


def tighten_paragraph(paragraph: str) -> str:
    """Further shorten a paragraph if the abstract is still too long."""
    sentences = split_sentences(paragraph)
    if not sentences:
        return paragraph.strip()
    if len(sentences) == 1:
        return sentences[0]
    return sentences[0] + sentences[-1]


def sentence_has_signal(sentence: str) -> bool:
    """Heuristic signal sentence detection."""
    tokens = [
        "结果",
        "表明",
        "因此",
        "最终",
        "推荐",
        "回退",
        "RMSE",
        "MAE",
        "R2",
        "R²",
        "Brier",
        "Bootstrap",
        "LSBoost",
        "PLS",
        "Firth",
        "BMI",
        "风险",
        "不能",
        "未能",
        "区分",
        "诊断",
    ]
    if any(token in sentence for token in tokens):
        return True
    return bool(re.search(r"\d|%|<|>|=|±|\\", sentence))


def split_sentences(text: str) -> list[str]:
    """Split Chinese text into sentence-like pieces."""
    parts = re.split(r"(?<=[。！？；])", text)
    return [part.strip() for part in parts if part.strip()]


def split_paragraphs(lines: list[str]) -> list[str]:
    """Combine consecutive non-empty lines into paragraphs."""
    paragraphs: list[str] = []
    buffer: list[str] = []
    for line in lines:
        if not line.strip():
            if buffer:
                paragraphs.append("\n".join(buffer).strip())
                buffer = []
            continue
        buffer.append(line.rstrip())
    if buffer:
        paragraphs.append("\n".join(buffer).strip())
    return paragraphs


def merge_image_blocks(markdown: str, work_dir: Path) -> str:
    """Merge adjacent image blocks into composite figures."""
    lines = markdown.splitlines()
    result: list[str] = []
    block: list[tuple[str, str]] = []
    blank_after_block = 0
    composite_index = 1

    def flush_block() -> None:
        nonlocal composite_index, blank_after_block
        if not block:
            return
        if len(block) == 1:
            alt, rel = block[0]
            result.append(f"![{alt}]({rel})")
        else:
            composite_rel = build_composite_figure(
                [work_dir / rel for _, rel in block],
                [alt for alt, _ in block],
                work_dir,
                composite_index,
            )
            composite_index += 1
            caption = "；".join(alt for alt, _ in block if alt) or "复合图"
            result.append(f"![{caption}]({composite_rel.as_posix()})")
        block.clear()
        blank_after_block = 0

    for line in lines:
        match = IMAGE_BLOCK_RE.match(line)
        if match:
            block.append((match.group("alt").strip(), match.group("src").strip()))
            continue

        if not line.strip():
            if block:
                blank_after_block += 1
                continue
            result.append("")
            continue

        if block:
            flush_block()
        result.append(line)

    if block:
        flush_block()
    return "\n".join(result)


def build_composite_figure(
    image_paths: list[Path],
    labels: list[str],
    work_dir: Path,
    index: int,
) -> Path:
    """Build a composite image from a small group of related figures."""
    available = [(path, label) for path, label in zip(image_paths, labels) if path.exists()]
    if not available:
        return image_paths[0]

    composite_dir = work_dir / "paper_composites"
    composite_dir.mkdir(parents=True, exist_ok=True)
    composite_path = composite_dir / f"composite_{index:03d}.png"

    images = [Image.open(path).convert("RGB") for path, _ in available]
    try:
        cols = 2 if len(images) > 1 else 1
        rows = math.ceil(len(images) / cols)
        cell_w = 920
        cell_h = 620
        caption_h = 72
        gap = 24
        canvas_w = cols * cell_w + (cols + 1) * gap
        canvas_h = rows * (cell_h + caption_h) + (rows + 1) * gap
        canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
        draw = ImageDraw.Draw(canvas)
        caption_font = load_font(work_dir, 22)
        title_font = load_font(work_dir, 26)

        for idx, (image, (_, label)) in enumerate(zip(images, available)):
            row = idx // cols
            col = idx % cols
            x0 = gap + col * (cell_w + gap)
            y0 = gap + row * (cell_h + caption_h + gap)
            tile = Image.new("RGB", (cell_w, cell_h + caption_h), "white")
            fitted = ImageOps.contain(image, (cell_w - 40, cell_h - 40))
            tile.paste(
                fitted,
                ((cell_w - fitted.width) // 2, 12 + (cell_h - fitted.height) // 2),
            )
            tile_draw = ImageDraw.Draw(tile)
            caption = label or f"图{idx + 1}"
            bbox = tile_draw.textbbox((0, 0), caption, font=caption_font)
            caption_w = bbox[2] - bbox[0]
            tile_draw.text(
                ((cell_w - caption_w) / 2, cell_h + 12),
                caption,
                fill="#333333",
                font=caption_font,
            )
            canvas.paste(tile, (x0, y0))
            draw.rounded_rectangle(
                [x0, y0, x0 + cell_w, y0 + cell_h + caption_h],
                radius=22,
                outline="#8A8A8A",
                width=2,
            )

        summary = " / ".join(label for label in labels if label.strip())
        if summary:
            draw.text(
                (canvas_w // 2, canvas_h - 18),
                summary[:120],
                fill="#666666",
                font=title_font,
                anchor="ms",
            )
        canvas.save(composite_path)
    finally:
        for image in images:
            image.close()

    return composite_path.relative_to(work_dir)


def load_font(work_dir: Path, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Load a Chinese-capable font."""
    backend_root = Path(__file__).resolve().parents[2]
    candidates = [
        work_dir / "simhei.ttf",
        backend_root / "fonts" / "simhei.ttf",
        backend_root / "fonts" / "SimHei.ttf",
        backend_root / "fonts" / "msyh.ttc",
        Path("C:/Windows/Fonts/simhei.ttf"),
        # macOS：PingFang.ttc 受 SIP 保护无法被 PIL 读取，选用实测可打开的字体
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
        Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                return ImageFont.truetype(str(candidate), size)
            except OSError:
                continue
    return ImageFont.load_default()


def append_source_code_appendix(markdown: str, work_dir: Path) -> str:
    """Append a compact source-code appendix."""
    if re.search(r"^\s*#\s*附录", markdown, flags=re.M) and "源代码" in markdown:
        return markdown

    candidates = collect_code_candidates(work_dir)
    if not candidates:
        return markdown

    appendix_lines = ["", "# 附录：源代码"]
    appendix_lines.append("本附录列出核心实现文件，便于复核建模流程与参数设定。")

    for idx, path in enumerate(candidates[:5], start=1):
        appendix_lines.append(f"## A.{idx} {path.name}")
        appendix_lines.append(f"文件作用：{describe_code_file(path.name)}")
        appendix_lines.append("")
        appendix_lines.append(f"```{language_for(path)}")
        appendix_lines.append(read_code_excerpt(path))
        appendix_lines.append("```")
        appendix_lines.append("")

    return markdown.rstrip() + "\n" + "\n".join(appendix_lines).rstrip() + "\n"


def collect_code_candidates(work_dir: Path) -> list[Path]:
    """Pick the most relevant code files for the appendix."""
    score_tokens = ("run", "fit", "predict", "policy", "search", "bootstrap", "stable", "core")
    candidates: list[tuple[int, Path]] = []
    for path in work_dir.iterdir():
        if not path.is_file() or path.suffix.lower() not in {".m", ".py"}:
            continue
        name = path.name.lower()
        score = 0
        if path.name == "build_polished_paper.py":
            score -= 100
        if name.startswith("q2"):
            score += 90
        if name.startswith("ques1") or name.startswith("ques2") or name.startswith("ques3") or name.startswith("ques4"):
            score += 80
        score += sum(15 for token in score_tokens if token in name)
        score += min(path.stat().st_size // 20000, 20)
        candidates.append((score, path))

    candidates.sort(key=lambda item: (-item[0], item[1].name))
    return [path for _, path in candidates]


def describe_code_file(filename: str) -> str:
    """Give a short human-readable description for a code file."""
    name = filename.lower()
    if "fit" in name:
        return "模型拟合与参数估计"
    if "predict" in name:
        return "结果预测与推断"
    if "bootstrap" in name:
        return "Bootstrap 稳健性评估"
    if "policy" in name or "search" in name:
        return "策略搜索与决策选择"
    if "stable" in name or "core" in name:
        return "核心求解流程"
    return "核心实现文件"


def language_for(path: Path) -> str:
    """Map file suffix to code fence language."""
    if path.suffix.lower() == ".m":
        return "matlab"
    return "python"


def read_code_excerpt(path: Path, max_chars: int = 4000) -> str:
    """Read a bounded excerpt of a code file for the appendix."""
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + "\n..."


def ensure_blank_lines_around_headings(markdown: str) -> str:
    """Keep blank lines around headings for stable Markdown parsing."""
    lines = markdown.splitlines()
    normalized: list[str] = []
    for idx, line in enumerate(lines):
        stripped = line.rstrip()
        if HEADING_RE.match(stripped):
            if normalized and normalized[-1].strip():
                normalized.append("")
            normalized.append(stripped)
            next_line = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
            if next_line and not HEADING_RE.match(next_line):
                normalized.append("")
        else:
            normalized.append(line)
    return "\n".join(normalized)


def build_reference_docx(work_dir: Path) -> Path:
    """Create a reference DOCX used by pandoc as a style template."""
    reference_path = work_dir / "paper_reference.docx"
    doc = Document()
    configure_document(doc)
    doc.save(reference_path)
    return reference_path


def configure_document(doc: Document) -> None:
    """Configure page size, margins, fonts and footer."""
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    set_style_font(styles["Normal"], east="SimSun", latin="Times New Roman", size=12)
    normal = styles["Normal"].paragraph_format
    normal.first_line_indent = Pt(24)
    normal.line_spacing = 1.32
    normal.space_before = Pt(0)
    normal.space_after = Pt(0)
    normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    style_specs = [
        ("Title", "SimHei", 17.5, 42, 18, WD_ALIGN_PARAGRAPH.CENTER, True),
        ("Heading 1", "SimHei", 15, 12, 8, WD_ALIGN_PARAGRAPH.CENTER, True),
        ("Heading 2", "SimHei", 13.5, 8, 5, WD_ALIGN_PARAGRAPH.LEFT, True),
        ("Heading 3", "SimHei", 12.5, 6, 3, WD_ALIGN_PARAGRAPH.LEFT, True),
        ("Caption", "SimSun", 10.5, 6, 4, WD_ALIGN_PARAGRAPH.CENTER, False),
    ]
    for name, east, size, before, after, align, bold in style_specs:
        style = styles[name]
        set_style_font(style, east=east, latin="Times New Roman", size=size, bold=bold)
        p = style.paragraph_format
        p.first_line_indent = Pt(0)
        p.line_spacing = 1.25
        p.space_before = Pt(before)
        p.space_after = Pt(after)
        p.alignment = align
        p.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Pt(0)
    add_page_field(footer)
    for run in footer.runs:
        set_run_font(run, east="SimSun", latin="Times New Roman", size=10.5)


def add_page_field(paragraph) -> None:
    """Insert a PAGE field in the footer."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_style_font(style, east="SimSun", latin="Times New Roman", size=12, bold=False):
    """Set font for a Word style."""
    font = style.font
    font.name = latin
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string("000000")
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_run_font(run, east="SimSun", latin="Times New Roman", size=12, bold=None):
    """Set font for a Word run."""
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def convert_markdown_to_docx(
    markdown: str,
    output_path: Path,
    reference_docx: Path,
    resource_path: Path,
) -> None:
    """Convert Markdown to DOCX using pandoc."""
    import pypandoc  # type: ignore[import-unresolved]

    pypandoc.convert_text(
        markdown,
        to="docx",
        format="markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables+raw_html",
        outputfile=str(output_path),
        extra_args=[
            f"--reference-doc={reference_docx}",
            f"--resource-path={resource_path}",
            "--standalone",
            "--wrap=none",
        ],
    )


def convert_markdown_to_pdf(
    markdown: str,
    output_path: Path,
    resource_path: Path,
) -> None:
    """Convert Markdown to PDF as a best-effort deliverable."""
    import pypandoc  # type: ignore[import-unresolved]

    header_path = resource_path / "paper_pdf_header.tex"
    header_path.write_text(build_pdf_header(resource_path), encoding="utf-8")

    try:
        pypandoc.convert_text(
            markdown,
            to="pdf",
            format="markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables+raw_html",
            outputfile=str(output_path),
            extra_args=[
                f"--resource-path={resource_path}",
                f"--include-in-header={header_path}",
                "--standalone",
                "--wrap=none",
                "--pdf-engine=xelatex",
            ],
        )
        return
    except Exception as exc:
        logger.warning("pandoc pdf export failed, falling back to HTML/WeasyPrint: %s", exc)

    html = pypandoc.convert_text(
        markdown,
        to="html5",
        format="markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables+raw_html",
        extra_args=[
            f"--resource-path={resource_path}",
            "--standalone",
            "--wrap=none",
            "--mathml",
        ],
    )

    from weasyprint import CSS, HTML  # type: ignore[import-unresolved]

    html_path = resource_path / "paper_render.html"
    html_path.write_text(html, encoding="utf-8")
    HTML(string=html, base_url=str(resource_path)).write_pdf(
        str(output_path),
        stylesheets=[CSS(string=build_weasyprint_css())],
    )


def build_weasyprint_css() -> str:
    """Build a lightweight CSS stylesheet for HTML-to-PDF fallback."""
    return """
@page {
  size: A4;
  margin: 2.6cm 2.7cm 2.2cm 2.7cm;
  @bottom-center {
    content: counter(page);
    font-family: "Times New Roman", "SimSun";
    font-size: 10.5pt;
  }
}

body {
  font-family: "Times New Roman", "SimSun";
  font-size: 12pt;
  line-height: 1.32;
  text-align: justify;
  color: #000;
}

h1, h2, h3, h4, h5, h6 {
  page-break-after: avoid;
  page-break-inside: avoid;
}

h1 {
  font-size: 17.5pt;
  text-align: center;
  font-weight: 700;
}

h2 {
  font-size: 15pt;
  text-align: center;
  font-weight: 700;
}

h3 {
  font-size: 13.5pt;
  font-weight: 700;
}

p {
  margin: 0 0 0.45em 0;
}

img {
  max-width: 100%;
  height: auto;
}

table {
  border-collapse: collapse;
  width: 100%;
}

th, td {
  border: 1px solid #444;
  padding: 0.22em 0.35em;
}

code, pre {
  white-space: pre-wrap;
  word-break: break-word;
}
"""


def build_pdf_header(
    resource_path: Path,
    comp_template: CompTemplate = CompTemplate.CHINA,
) -> str:
    """Create a portable XeLaTeX header for the selected competition."""
    paper = "a4paper" if comp_template == CompTemplate.CHINA else "letterpaper"
    if (resource_path / "simhei.ttf").is_file():
        cjk_font = r"""\setCJKmainfont[
  Path=./,
  BoldFont=simhei.ttf
]{simhei.ttf}"""
    else:
        # Docker 使用系统字体；Fandol 兼容未安装 Noto 的 TeX 发行版。
        cjk_font = r"""\IfFontExistsTF{Noto Sans CJK SC}{
  \setCJKmainfont{Noto Sans CJK SC}
}{
  \setCJKmainfont[BoldFont=FandolSong-Bold]{FandolSong-Regular}
}"""
    return f"""
\\usepackage[{paper},top=2.4cm,bottom=2.3cm,left=2.5cm,right=2.5cm]{{geometry}}
\\usepackage{{fontspec}}
\\usepackage{{xeCJK}}
\\usepackage{{amsmath,amssymb}}
\\usepackage{{booktabs,longtable,array,graphicx,float}}
\\IfFontExistsTF{{Times New Roman}}{{\\setmainfont{{Times New Roman}}}}{{\\setmainfont{{TeX Gyre Termes}}}}
{cjk_font}
\\IfFontExistsTF{{Arial}}{{\\setsansfont{{Arial}}}}{{\\setsansfont{{TeX Gyre Heros}}}}
\\IfFontExistsTF{{Consolas}}{{\\setmonofont{{Consolas}}}}{{\\setmonofont{{Latin Modern Mono}}}}
\\XeTeXlinebreaklocale "zh"
\\XeTeXlinebreakskip = 0pt plus 1pt
\\setlength{{\\emergencystretch}}{{3em}}
\\setlength{{\\parindent}}{{2em}}
\\setlength{{\\parskip}}{{0.25em}}
"""
